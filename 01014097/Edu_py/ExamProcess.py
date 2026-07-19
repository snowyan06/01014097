import os
import re
import logging
from typing import Dict, Optional, Tuple
from http import HTTPStatus
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from dashscope import Application
import datetime
# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename='exam_processor.log'
)


def get_current_timestamp():
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
class ExamProcessor:
    def __init__(self):
        # 从环境变量获取输出目录，默认为当前目录下的output
        self.output_dir = os.getenv('OUTPUT_DIR', './output')
        os.makedirs(self.output_dir, exist_ok=True)
        self.api_key = os.getenv('OTHER_API_KEY', 'sk-10eea048e8244237b67db0e87a590cd3')
        self.app_id = os.getenv('APP_ID', 'afd46bbff4ce490388661f3f02fd283f')

    def parse_file(self, file_path: str) -> Optional[str]:
        """解析文件内容"""
        try:
            if file_path.endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return "\n".join([page.extract_text() for page in reader.pages])
            elif file_path.endswith('.docx'):
                doc = Document(file_path)
                return "\n".join([p.text for p in doc.paragraphs])
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                logging.error(f"Unsupported format: {file_path}")
                return None
        except Exception as e:
            logging.error(f"Failed to parse file: {str(e)}")
            return None

    def extract_questions(self, text: str) -> str:
        """提取试题部分（保留题干编号和内容）"""
        # 分割题目和答案部分
        questions = re.split(r'参考答案|答案部分|评分标准', text, flags=re.IGNORECASE)[0]

        # 提取带编号的题目
        question_pattern = re.compile(r'(\d+[\.、．].*?(?=\n\d+[\.、．]|\Z))', re.DOTALL)
        questions_list = question_pattern.findall(questions)

        # 重新组合确保格式统一
        return "\n\n".join([q.strip() for q in questions_list if q.strip()])

    def get_standard_answers(self, questions: str) -> Optional[str]:
        """第一次调用：获取带题干的标准答案"""
        prompt = f"""请根据以下试题内容，为每道题提供标准答案，要求保留原题干：

试题内容：
{questions}

要求：
1. 保留原题干内容
2. 在每个题干下方添加【答案】部分
3. 选择题标明正确选项
4. 主观题列出关键得分点

请严格按以下格式返回：
【标准答案】

1. 第一题题干内容...
【答案】第一题答案内容...

2. 第二题题干内容...
【答案】第二题答案内容...
..."""

        try:
            response = Application.call(
                api_key=self.api_key,
                app_id=self.app_id,
                prompt=prompt
            )
            if response.status_code == HTTPStatus.OK:
                return response.output.text
            logging.error(f"获取答案失败: {response.message}")
            return None
        except Exception as e:
            logging.error(f"API调用异常: {str(e)}")
            return None

    def generate_grading_rubric(self, questions: str, answers: str) -> Optional[str]:
        """第二次调用：生成带题干的评分标准"""
        prompt = f"""请根据以下试题和答案，生成包含原题干的评分标准：

试题内容：
{questions}

标准答案：
{answers}

要求：
1. 保留原题干内容
2. 在每个题干下方添加【评分标准】部分
3. 说明每题总分和得分点
4. 主观题需详细说明评分细则

请严格按以下格式返回：
【评分标准】

1. 第一题题干内容...
【评分标准】
- 本题总分：X分
- 得分点1：...（X分）
- 得分点2：...（X分）

2. 第二题题干内容...
【评分标准】
..."""

        try:
            response = Application.call(
                api_key=self.api_key,
                app_id=self.app_id,
                prompt=prompt
            )
            if response.status_code == HTTPStatus.OK:
                return response.output.text
            logging.error(f"生成评分标准失败: {response.message}")
            return None
        except Exception as e:
            logging.error(f"API调用异常: {str(e)}")
            return None

    def save_to_file(self, content: str, filename: str) -> bool:
        """保存内容到docx文件，优化排版和格式"""
        try:
            os.makedirs('output', exist_ok=True)

            # 创建Word文档
            doc = Document()

            # 设置文档默认样式
            self._set_document_styles(doc)

            # 预处理内容：分割段落并去除空行
            sections = [s for s in content.split('\n\n') if s.strip()]

            current_question = None

            for section in sections:
                # 处理标题部分（如【标准答案】）
                if section.strip().startswith('【') and section.strip().endswith('】'):
                    self._add_section_title(doc, section.strip())
                    continue

                # 处理题目和内容
                lines = [line.strip() for line in section.split('\n') if line.strip()]
                if not lines:
                    continue

                # 处理题目编号（如1. 或1、）
                if re.match(r'^\d+[\.、．]', lines[0]):
                    current_question = lines[0]
                    # 添加题目段落
                    p = doc.add_paragraph(style='QuestionStyle')
                    p.add_run(current_question).bold = True

                    # 添加题目内容（如果有）
                    if len(lines) > 1:
                        content = ' '.join(lines[1:])
                        p = doc.add_paragraph(style='ContentStyle')
                        p.add_run(content)

                # 处理答案或评分标准
                elif any(lines[0].startswith(tag) for tag in ['【答案】', '【评分标准】']):
                    tag_type = lines[0][:lines[0].find('】') + 1]
                    content = ' '.join([lines[0][len(tag_type):]] + lines[1:])

                    # 添加标签段落
                    p = doc.add_paragraph(style='TagStyle')
                    run = p.add_run(tag_type)
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)  # 红色

                    # 添加内容
                    p.add_run(content)

                    # 如果是评分标准，特殊处理
                    if tag_type == '【评分标准】':
                        self._format_rubric_content(doc, content)
                else:
                    # 普通内容段落
                    p = doc.add_paragraph(style='ContentStyle')
                    p.add_run(' '.join(lines))

            # 设置页边距（单位：厘米）
            sections = doc.sections
            for section in sections:
                section.top_margin = Pt(28.35)  # 1厘米
                section.bottom_margin = Pt(28.35)
                section.left_margin = Pt(28.35)
                section.right_margin = Pt(28.35)

            # 保存文档
            output_path = filename if filename.endswith('.docx') else filename.replace('.txt', '.docx')
            doc.save(output_path)
            logging.info(f"文件已保存: {output_path}")
            return True
        except Exception as e:
            logging.error(f"保存文件时出错: {str(e)}")
            return False

    def _set_document_styles(self, doc):
        """设置文档样式"""
        # 设置正文样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 添加题目样式
        if 'QuestionStyle' not in doc.styles:
            question_style = doc.styles.add_style('QuestionStyle', 1)  # 1表示段落样式
            question_font = question_style.font
            question_font.name = '黑体'
            question_font.size = Pt(14)
            question_font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            question_style.paragraph_format.space_after = Pt(6)
            question_style.paragraph_format.line_spacing = 1.5

        # 添加内容样式
        if 'ContentStyle' not in doc.styles:
            content_style = doc.styles.add_style('ContentStyle', 1)
            content_font = content_style.font
            content_font.name = '宋体'
            content_font.size = Pt(12)
            content_font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            content_style.paragraph_format.space_before = Pt(0)
            content_style.paragraph_format.space_after = Pt(12)
            content_style.paragraph_format.line_spacing = 1.5

        # 添加标签样式（答案/评分标准）
        if 'TagStyle' not in doc.styles:
            tag_style = doc.styles.add_style('TagStyle', 1)
            tag_font = tag_style.font
            tag_font.name = '楷体'
            tag_font.size = Pt(12)
            tag_font._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            tag_style.paragraph_format.space_before = Pt(6)
            tag_style.paragraph_format.space_after = Pt(6)
            tag_style.paragraph_format.line_spacing = 1.5

    def _add_section_title(self, doc, title_text):
        """添加章节标题"""
        p = doc.add_paragraph(style='Heading 1')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        doc.add_paragraph()  # 添加空行

    def _format_rubric_content(self, doc, content):
        """特殊处理评分标准内容"""
        # 分割得分点
        points = [p.strip() for p in content.split('-') if p.strip()]

        for point in points:
            # 处理总分行
            if '本题总分' in point:
                p = doc.add_paragraph(style='TagStyle')
                run = p.add_run(f"- {point}")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
            else:
                p = doc.add_paragraph(style='ContentStyle')
                p.add_run(f"- {point}")
                p.paragraph_format.left_indent = Pt(18)  # 缩进

    def process_exam(self, file_path: str) -> Tuple[bool, str, Optional[str], Optional[str], Optional[int], Optional[int]]:
        """处理试卷完整流程"""
        # 1. 解析文件
        text = self.parse_file(file_path)
        if not text:
            return False, "文件解析失败", None, None,None, None

        # 2. 提取试题部分（带题干）
        questions = self.extract_questions(text)
        if not questions:
            return False, "试题提取失败", None, None,None, None
        from datetime import datetime
        # 生成时间戳
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 构建文件名
        answer_file = os.path.join(self.output_dir, f"标准答案_{timestamp}.docx")
        rubric_file = os.path.join(self.output_dir, f"评分标准_{timestamp}.docx")

        # 3. 第一次调用获取带题干的标准答案
        answers = self.get_standard_answers(questions)
        if not answers:
            return False, "获取答案失败", None, None,None, None
        self.save_to_file(answers, answer_file)

        # 4. 第二次调用生成带题干的评分标准
        rubric = self.generate_grading_rubric(questions, answers)
        if not rubric:
            return False, "生成评分标准失败", None, None,None, None
        self.save_to_file(rubric, rubric_file)
        answer_size = os.path.getsize(answer_file)
        rubric_size = os.path.getsize(rubric_file)
        answer_filename = os.path.basename(answer_file)
        rubric_filename = os.path.basename(rubric_file)

        return True, "处理成功", answer_filename, rubric_filename ,answer_size,rubric_size


if __name__ == "__main__":
    processor = ExamProcessor()
    result, message = processor.process_exam("exam.pdf")  # 替换为您的文件路径
    print(f"处理结果: {message}")
    print("生成的答案和评分标准已保存到output目录")