"""
试卷分析 Agent
迁移自 exam_analysis.py，负责：
- 读取 Excel 考试成绩数据
- 调用 LLM 生成分析报告内容
- 生成 Word 试卷分析表 + 成绩分布图
"""

import os
import re
import logging
import numpy as np
import openpyxl
import matplotlib.pyplot as plt
from datetime import datetime
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from agents.base_agent import BaseAgent
from config.settings import settings

logger = logging.getLogger(__name__)


class ExamAnalysisAgent(BaseAgent):
    """试卷分析 Agent"""

    agent_name: str = "exam_analysis"
    agent_description: str = "读取 Excel 成绩数据，生成试卷分析报告 Word 文档"

    def __init__(self, adapter=None, app_id=None):
        super().__init__(adapter, app_id or settings.APP_ID)
        self.output_dir = settings.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)

    # ================================================================
    # 数据读取与统计
    # ================================================================

    def _read_excel_data(self, file_path: str) -> Optional[Dict]:
        """读取 Excel 文件并提取考试数据"""
        try:
            wb = openpyxl.load_workbook(file_path)
            sheet = wb.active

            required_cells = {
                'B1': '学年', 'B2': '学期',
                'B3': '课程名称', 'B4': '命题人',
                'B5': '考试班级', 'B6': '考试日期'
            }

            data = {}
            for cell_ref, field in required_cells.items():
                value = sheet[cell_ref].value
                if not value:
                    raise ValueError(f"缺少必要字段: {field}")
                data[field] = str(value).strip()

            data['scores'] = []
            objectives = {}

            for row in sheet.iter_rows(min_row=8, max_row=sheet.max_row):
                if row[0].value is not None:
                    try:
                        data['scores'].append(float(row[0].value))
                    except (ValueError, TypeError):
                        continue

                if len(data.get('course_objectives', [])) < 4 and row[1].value:
                    try:
                        obj_str = str(row[1].value)
                        if "课程目标" in obj_str:
                            obj_num = int(obj_str.split("课程目标")[1].split("：")[0])
                            objectives[obj_num] = {
                                'index': obj_num,
                                'requirement': str(row[2].value).strip(),
                                'full_score': float(row[3].value),
                                'absolute_score': 0.0
                            }
                    except (IndexError, ValueError, AttributeError):
                        continue

            if not data['scores']:
                raise ValueError("未找到有效成绩数据")

            avg_score = np.mean(data['scores'])
            for obj in objectives.values():
                obj['absolute_score'] = round(avg_score * (obj['full_score'] / 100), 2)
                obj['evaluation'] = round(obj['absolute_score'] / obj['full_score'], 3)
                obj['analysis'] = f"课程目标{obj['index']}达成度为{obj['evaluation']}，详细分析待补充。"

            data['course_objectives'] = list(objectives.values())
            data['problems'] = [
                "学生基础知识掌握不牢固",
                "综合分析能力有待提高",
                "解题思路不够清晰",
                "答题规范性不足"
            ]
            data['improvement_measures'] = [
                "加强基础知识的讲解和练习",
                "增加案例分析和讨论环节",
                "提供更多解题思路训练",
                "规范答题格式要求"
            ]

            return {
                'academic_year': data['学年'],
                'semester': data['学期'],
                'course_name': data['课程名称'],
                'teacher_name': data['命题人'],
                'class_name': data['考试班级'],
                'exam_date': data['考试日期'],
                'scores': data['scores'],
                'course_objectives': data['course_objectives'],
                'problems': data['problems'],
                'improvement_measures': data['improvement_measures']
            }

        except Exception as e:
            logger.error("Excel 文件解析失败: %s", e, exc_info=True)
            return None

    @staticmethod
    def _calculate_statistics(scores: List[int]) -> Dict:
        scores_array = np.array(scores)
        return {
            'average': round(float(np.mean(scores_array)), 2),
            'excellent_rate': round(float(np.sum(scores_array >= 90) / len(scores) * 100), 2),
            'good_rate': round(float(np.sum((scores_array >= 80) & (scores_array < 90)) / len(scores) * 100), 2),
            'fail_rate': round(float(np.sum(scores_array < 60) / len(scores) * 100), 2)
        }

    def _generate_score_distribution_chart(self, scores: List[int], filename: str) -> Optional[str]:
        try:
            plt.rcParams['font.sans-serif'] = ['SimHei']
            plt.rcParams['axes.unicode_minus'] = False

            fig, ax = plt.subplots(figsize=(8, 4))
            bins = [0, 60, 70, 80, 90, 100]
            labels = ['不及格', '60-69', '70-79', '80-89', '90-100']

            hist, _ = np.histogram(scores, bins=bins)
            bars = ax.bar(labels, hist, color='#4472C4')

            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width() / 2., height,
                        f'{int(height)}', ha='center', va='bottom')

            ax.set_title('成绩分布图', fontsize=12)
            ax.set_xlabel('分数段', fontsize=10)
            ax.set_ylabel('人数', fontsize=10)

            chart_path = os.path.join(self.output_dir, filename)
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            return chart_path
        except Exception as e:
            logger.error("生成成绩分布图失败: %s", e)
            return None

    # ================================================================
    # LLM 分析
    # ================================================================

    def _generate_analysis_content(self, exam_data: Dict) -> Optional[Dict]:
        """调用 LLM 生成试卷分析内容"""
        stats = self._calculate_statistics(exam_data['scores'])
        prompt = f"""请根据以下考试数据生成专业的试卷分析报告：

课程名称：{exam_data['course_name']}
考试班级：{exam_data['class_name']}
平均成绩：{stats['average']}分

课程目标达成情况：
{chr(10).join([f"课程目标{obj['index']}：满分{obj['full_score']}分，平均得分{obj['absolute_score']}分" for obj in exam_data['course_objectives']])}

请按照以下结构生成分析报告：
1. 对每个课程目标的达成度分析和教学建议
2. 学生答题中出现的主要问题（至少列出4个）
3. 改进教学的具体措施（至少列出4条）

要求：
- 使用专业的教育评估术语
- 每个课程目标分析不少于100字
- 问题分析要具体
- 改进措施要具有可操作性"""

        response = self.call_llm(prompt)
        if not response.success:
            logger.warning("LLM 返回失败，使用默认分析内容: %s", response.error_message)
            return exam_data

        analysis = response.content
        try:
            objectives_section = re.search(r'1\..*?(课程目标[\s\S]*?)(?=2\.|\Z)', analysis, re.IGNORECASE)
            if objectives_section:
                objectives_text = objectives_section.group(1)
                for obj in exam_data['course_objectives']:
                    obj_pattern = re.compile(
                        fr'课程目标{obj["index"]}[：:](.*?)(?=课程目标|\Z)', re.DOTALL
                    )
                    match = obj_pattern.search(objectives_text)
                    if match:
                        obj['analysis'] = match.group(1).strip()

            problems_section = re.search(r'2\..*?主要问题([\s\S]*?)(?=3\.|\Z)', analysis, re.IGNORECASE)
            if problems_section:
                problems_text = problems_section.group(1)
                problems = [p.strip() for p in re.findall(r'[•\-]\s*(.*?)(?=\n|$)', problems_text) if p.strip()]
                if problems:
                    exam_data['problems'] = problems[:4] if len(problems) >= 4 else problems + [""] * (4 - len(problems))

            measures_section = re.search(r'3\..*?改进措施([\s\S]*?)(?=\Z)', analysis, re.IGNORECASE)
            if measures_section:
                measures_text = measures_section.group(1)
                measures = [m.strip() for m in re.findall(r'[•\-]\s*(.*?)(?=\n|$)', measures_text) if m.strip()]
                if measures:
                    exam_data['improvement_measures'] = measures[:4] if len(measures) >= 4 else measures + [""] * (4 - len(measures))

            return exam_data
        except Exception as e:
            logger.error("解析 LLM 返回内容失败: %s", e)
            return exam_data

    # ================================================================
    # Word 报告生成
    # ================================================================

    def _generate_word_report(self, exam_data: Dict) -> Optional[str]:
        try:
            required_fields = ['academic_year', 'semester', 'course_name', 'teacher_name',
                               'class_name', 'exam_date', 'scores', 'course_objectives',
                               'problems', 'improvement_measures']
            for field in required_fields:
                if field not in exam_data:
                    raise ValueError(f"缺少必要字段: {field}")

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)

            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.add_run("XX大学信息工程学院（人工智能学院）\n\n试 卷 分 析 表\n")
            title_run.font.size = Pt(16)
            title_run.bold = True

            semester = doc.add_paragraph()
            semester.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            semester.add_run(f"（{exam_data['academic_year']}学年第 {exam_data['semester']} 学期）\n")

            basic_info = doc.add_table(rows=2, cols=4)
            basic_info.style = 'Table Grid'
            for row in basic_info.rows:
                for cell in row.cells:
                    cell.width = Inches(2.2)

            cells = basic_info.rows[0].cells
            cells[0].text = "课程名称"
            cells[1].text = exam_data['course_name']
            cells[2].text = "命题人"
            cells[3].text = exam_data['teacher_name']

            cells = basic_info.rows[1].cells
            cells[0].text = "考试班级"
            cells[1].text = exam_data['class_name']
            cells[2].text = "考试日期"
            cells[3].text = exam_data['exam_date']

            doc.add_paragraph("\n1、试卷成绩分析", style='Heading 1')
            doc.add_paragraph("（1）试卷成绩分布", style='Heading 2')

            stats = self._calculate_statistics(exam_data['scores'])
            excellent_good_rate = stats['excellent_rate'] + stats['good_rate']
            doc.add_paragraph(
                f"成绩分布较合理，平均成绩{stats['average']}分，"
                f"优良率{excellent_good_rate}%，不及格率{stats['fail_rate']}%。"
            )

            chart_path = self._generate_score_distribution_chart(exam_data['scores'], "score_distribution.png")
            if chart_path:
                doc.add_paragraph("成绩分布图如下：")
                doc.add_picture(chart_path, width=Inches(4.0))

            doc.add_paragraph("\n2、试卷对课程目标的达成度分析", style='Heading 1')
            for obj in exam_data['course_objectives']:
                analysis_text = obj.get('analysis', f"课程目标{obj['index']}达成度为{obj['evaluation']}，详细分析待补充。")
                doc.add_paragraph(
                    f"（{obj['index']}）课程目标 {obj['index']}（支撑毕业要求 {obj['requirement']}）"
                    f"达成度为 {obj['evaluation']}，{analysis_text}"
                )

            doc.add_paragraph("\n3、学生答题主要问题分析：", style='Heading 1')
            for i, problem in enumerate(exam_data.get('problems', []), 1):
                doc.add_paragraph(f"（{i}）{problem}")

            doc.add_paragraph("\n4、存在问题及改进措施", style='Heading 1')
            for i, measure in enumerate(exam_data.get('improvement_measures', []), 1):
                doc.add_paragraph(f"（{i}）{measure}")

            doc.add_paragraph("\n\n签名：\t\t\t\t年\t月\t日")

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(self.output_dir, f"{exam_data['course_name']}_{timestamp}_试卷分析表.docx")
            doc.save(output_path)
            return output_path
        except Exception as e:
            logger.error("生成 Word 报告失败: %s", e, exc_info=True)
            return None

    # ================================================================
    # 主处理方法
    # ================================================================

    def process(self, **kwargs) -> Dict[str, Any]:
        """
        处理 Excel 文件并生成试卷分析表

        :param file_path: 上传的 Excel 文件路径
        """
        file_path = kwargs.get("file_path")
        if not file_path:
            return self._build_error_response("缺少 file_path 参数")

        try:
            exam_data = self._read_excel_data(file_path)
            if not exam_data:
                return self._build_error_response("无法读取 Excel 数据")

            analyzed_data = self._generate_analysis_content(exam_data)
            if not analyzed_data:
                return self._build_error_response("无法生成分析内容")

            result = self._generate_word_report(analyzed_data)
            if not result:
                return self._build_error_response("无法生成 Word 报告")

            report_filename = os.path.basename(result)
            report_size = os.path.getsize(result)

            return self._build_success_response(
                data={
                    "report_path": result,
                    "report_filename": report_filename,
                    "report_size": report_size,
                },
                message="试卷分析报告生成成功"
            )

        except Exception as e:
            logger.error("[Agent:exam_analysis] 异常: %s", e, exc_info=True)
            return self._build_error_response(str(e))
