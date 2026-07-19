"""
试卷处理 Agent
迁移自 ExamProcess.py，负责：
- 解析上传的试卷文件（PDF/DOCX/TXT）
- 调用 LLM 生成标准答案和评分标准
- 输出格式化 Word 文档
"""

import os
import re
import logging
import datetime
from typing import Dict, Any, Optional

import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from agents.base_agent import BaseAgent
from config.settings import settings

logger = logging.getLogger(__name__)


class ExamProcessAgent(BaseAgent):
    """试卷处理 Agent"""

    agent_name: str = "exam_process"
    agent_description: str = "解析试卷文件，生成标准答案和评分标准 Word 文档"

    def __init__(self, adapter=None, app_id=None):
        super().__init__(adapter, app_id or settings.APP_ID)
        self.output_dir = settings.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)

    # ================================================================
    # 文件解析
    # ================================================================

    @staticmethod
    def _parse_file(file_path: str) -> Optional[str]:
        """解析文件内容（支持 PDF / DOCX / TXT）"""
        try:
            if file_path.endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return "\n".join([page.extract_text() for page in reader.pages])
            elif file_path.endswith('.docx'):
                doc = Document(file_path)
                return "\n".join([p.text for p in doc.paragraphs])
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                logger.error("不支持的文件格式: %s", file_path)
                return None
        except Exception as e:
            logger.error("文件解析失败: %s", e)
            return None

    @staticmethod
    def _extract_questions(text: str) -> str:
        """提取试题部分"""
        questions = re.split(r'参考答案|答案部分|评分标准', text, flags=re.IGNORECASE)[0]
        question_pattern = re.compile(r'(\d+[\.、．].*?(?=\n\d+[\.、．]|\Z))', re.DOTALL)
        questions_list = question_pattern.findall(questions)
        return "\n\n".join([q.strip() for q in questions_list if q.strip()])

    # ================================================================
    # LLM 调用
    # ================================================================

    def _get_standard_answers(self, questions: str) -> Optional[str]:
        """第一次调用：获取带题干的标准答案"""
        prompt = f"""请根据以下试题内容，为每道题提供标准答案，要求保留原题干：

试题内容：
{questions}

要求：
1. 保留原题干内容
2. 在每个题干下方添加【答案】部分
3. 选择题标明正确选项
4. 主观题列出关键得分点

请严格按以下格式返回：
【标准答案】

1. 第一题题干内容...
【答案】第一题答案内容...

2. 第二题题干内容...
【答案】第二题答案内容...
..."""
        response = self.call_llm(prompt)
        if response.success:
            return response.content
        logger.error("获取答案失败: %s", response.error_message)
        return None

    def _generate_grading_rubric(self, questions: str, answers: str) -> Optional[str]:
        """第二次调用：生成带题干的评分标准"""
        prompt = f"""请根据以下试题和答案，生成包含原题干的评分标准：

试题内容：
{questions}

标准答案：
{answers}

要求：
1. 保留原题干内容
2. 在每个题干下方添加【评分标准】部分
3. 说明每题总分和得分点
4. 主观题需详细说明评分细则

请严格按以下格式返回：
【评分标准】

1. 第一题题干内容...
【评分标准】
- 本题总分：X分
- 得分点1：...（X分）
- 得分点2：...（X分）

2. 第二题题干内容...
【评分标准】
..."""
        response = self.call_llm(prompt)
        if response.success:
            return response.content
        logger.error("生成评分标准失败: %s", response.error_message)
        return None

    # ================================================================
    # Word 文档生成
    # ================================================================

    def _set_document_styles(self, doc):
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        if 'QuestionStyle' not in doc.styles:
            qs = doc.styles.add_style('QuestionStyle', 1)
            qs.font.name = '黑体'
            qs.font.size = Pt(14)
            qs.font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            qs.paragraph_format.space_after = Pt(6)
            qs.paragraph_format.line_spacing = 1.5

        if 'ContentStyle' not in doc.styles:
            cs = doc.styles.add_style('ContentStyle', 1)
            cs.font.name = '宋体'
            cs.font.size = Pt(12)
            cs.font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            cs.paragraph_format.space_before = Pt(0)
            cs.paragraph_format.space_after = Pt(12)
            cs.paragraph_format.line_spacing = 1.5

        if 'TagStyle' not in doc.styles:
            ts = doc.styles.add_style('TagStyle', 1)
            ts.font.name = '楷体'
            ts.font.size = Pt(12)
            ts.font._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            ts.paragraph_format.space_before = Pt(6)
            ts.paragraph_format.space_after = Pt(6)
            ts.paragraph_format.line_spacing = 1.5

    def _add_section_title(self, doc, title_text):
        p = doc.add_paragraph(style='Heading 1')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        doc.add_paragraph()

    def _format_rubric_content(self, doc, content):
        points = [p.strip() for p in content.split('-') if p.strip()]
        for point in points:
            if '本题总分' in point:
                p = doc.add_paragraph(style='TagStyle')
                run = p.add_run(f"- {point}")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 255)
            else:
                p = doc.add_paragraph(style='ContentStyle')
                p.add_run(f"- {point}")
                p.paragraph_format.left_indent = Pt(18)

    def _save_to_file(self, content: str, filename: str) -> bool:
        """保存内容到 docx 文件"""
        try:
            os.makedirs(self.output_dir, exist_ok=True)
            doc = Document()
            self._set_document_styles(doc)

            sections = [s for s in content.split('\n\n') if s.strip()]

            for section in sections:
                if section.strip().startswith('【') and section.strip().endswith('】'):
                    self._add_section_title(doc, section.strip())
                    continue

                lines = [line.strip() for line in section.split('\n') if line.strip()]
                if not lines:
                    continue

                if re.match(r'^\d+[\.、．]', lines[0]):
                    p = doc.add_paragraph(style='QuestionStyle')
                    p.add_run(lines[0]).bold = True
                    if len(lines) > 1:
                        p2 = doc.add_paragraph(style='ContentStyle')
                        p2.add_run(' '.join(lines[1:]))

                elif any(lines[0].startswith(tag) for tag in ['【答案】', '【评分标准】']):
                    tag_type = lines[0][:lines[0].find('】') + 1]
                    tag_content = ' '.join([lines[0][len(tag_type):]] + lines[1:])
                    p = doc.add_paragraph(style='TagStyle')
                    run = p.add_run(tag_type)
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    p.add_run(tag_content)
                    if tag_type == '【评分标准】':
                        self._format_rubric_content(doc, tag_content)
                else:
                    p = doc.add_paragraph(style='ContentStyle')
                    p.add_run(' '.join(lines))

            for section in doc.sections:
                section.top_margin = Pt(28.35)
                section.bottom_margin = Pt(28.35)
                section.left_margin = Pt(28.35)
                section.right_margin = Pt(28.35)

            output_path = filename if filename.endswith('.docx') else filename.replace('.txt', '.docx')
            doc.save(output_path)
            logger.info("文件已保存: %s", output_path)
            return True
        except Exception as e:
            logger.error("保存文件失败: %s", e)
            return False

    # ================================================================
    # 主处理方法
    # ================================================================

    def process(self, **kwargs) -> Dict[str, Any]:
        """
        处理试卷完整流程

        :param file_path: 上传的试卷文件路径
        """
        file_path = kwargs.get("file_path")
        if not file_path:
            return self._build_error_response("缺少 file_path 参数")

        text = self._parse_file(file_path)
        if not text:
            return self._build_error_response("文件解析失败")

        questions = self._extract_questions(text)
        if not questions:
            return self._build_error_response("试题提取失败")

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        answer_file = os.path.join(self.output_dir, f"标准答案_{timestamp}.docx")
        rubric_file = os.path.join(self.output_dir, f"评分标准_{timestamp}.docx")

        answers = self._get_standard_answers(questions)
        if not answers:
            return self._build_error_response("获取答案失败")
        self._save_to_file(answers, answer_file)

        rubric = self._generate_grading_rubric(questions, answers)
        if not rubric:
            return self._build_error_response("生成评分标准失败")
        self._save_to_file(rubric, rubric_file)

        answer_size = os.path.getsize(answer_file)
        rubric_size = os.path.getsize(rubric_file)
        answer_filename = os.path.basename(answer_file)
        rubric_filename = os.path.basename(rubric_file)

        return self._build_success_response(
            data={
                "answer_file": answer_filename,
                "answer_size": answer_size,
                "rubric_file": rubric_filename,
                "rubric_size": rubric_size,
            },
            message="处理成功"
        )
