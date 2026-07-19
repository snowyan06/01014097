import os
from docx import Document
import openpyxl
import json


class DocumentProcessor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.file_extension = os.path.splitext(file_path)[1].lower()

    def process(self):
        """主处理函数，根据文件类型调用对应处理方法"""
        if self.file_extension == '.docx':
            return self._process_word()
        elif self.file_extension == '.xlsx':
            return self._process_excel()
        elif self.file_extension == '.txt':
            return self._process_text()
        else:
            raise ValueError(f"不支持的文件类型: {self.file_extension}")

    def _process_word(self):
        """处理 Word 文档"""
        doc = Document(self.file_path)
        content = {
            'text': [],
            'tables': []
        }

        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content['text'].append(paragraph.text)

        # 提取表格数据
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append(table_data)

        return content

    def _process_excel(self):
        """处理 Excel 文件"""
        workbook = openpyxl.load_workbook(self.file_path)
        content = {}

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = []

            # 提取单元格数据
            for row in sheet.iter_rows(values_only=True):
                sheet_data.append(list(row))

            content[sheet_name] = sheet_data

        return content

    def _process_text(self):
        """处理纯文本文件"""
        with open(self.file_path, 'r', encoding='utf-8') as f:
            lines = [line.strip() for line in f.readlines() if line.strip()]

        return {
            'lines': lines,
            'line_count': len(lines),
            'file_info': {
                'file_name': os.path.basename(self.file_path),
                'file_size': os.path.getsize(self.file_path)
            }
        }


# 使用示例
if __name__ == "__main__":
    file_path = "example.xlsx"  # 替换为你的文件路径
    processor = DocumentProcessor(file_path)

    try:
        result = processor.process()
        print(result)
        # 将结果保存为 JSON 文件
        output_file = os.path.splitext(file_path)[0] + "_processed.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"处理完成，结果已保存至：{output_file}")
    except Exception as e:
        print(f"处理失败: {str(e)}")