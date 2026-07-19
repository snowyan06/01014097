import os
import re
from datetime import datetime

import openpyxl
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import numpy as np
import logging
from typing import List, Dict, Optional
from dashscope import Application

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename='exam_analysis.log'
)


class ExamAnalysisGenerator:
    def __init__(self):
        # 从环境变量获取输出目录
        self.output_dir = os.getenv('OUTPUT_DIR', './output')
        os.makedirs(self.output_dir, exist_ok=True)
        self.api_key = os.getenv('OTHER_API_KEY', 'sk-10eea048e8244237b67db0e87a590cd3')  # 替换为您的API Key
        self.app_id = os.getenv('APP_ID', 'afd46bbff4ce490388661f3f02fd283f')  # 替换为您的App ID

    def read_excel_data(self, file_path: str) -> Optional[Dict]:
        """读取Excel文件并提取考试数据"""
        try:
            wb = openpyxl.load_workbook(file_path)
            sheet = wb.active

            # 基础信息检查
            required_cells = {
                'B1': '学年', 'B2': '学期',
                'B3': '课程名称', 'B4': '命题人',
                'B5': '考试班级', 'B6': '考试日期'
            }

            data = {}
            for cell_ref, field in required_cells.items():
                value = sheet[cell_ref].value
                if not value:
                    raise ValueError(f"缺少必要字段: {field}")
                data[field] = str(value).strip()

            # 成绩数据读取
            data['scores'] = []
            data['course_objectives'] = []
            objectives = {}

            for row in sheet.iter_rows(min_row=8, max_row=sheet.max_row):
                # 处理成绩
                if row[0].value is not None:
                    try:
                        score = float(row[0].value)
                        data['scores'].append(score)
                    except (ValueError, TypeError):
                        continue

                # 处理课程目标（仅处理前4行）
                if len(data['course_objectives']) < 4 and row[1].value:
                    try:
                        obj_str = str(row[1].value)
                        if "课程目标" in obj_str:
                            obj_num = int(obj_str.split("课程目标")[1].split("：")[0])
                            objectives[obj_num] = {
                                'index': obj_num,
                                'requirement': str(row[2].value).strip(),
                                'full_score': float(row[3].value),
                                'absolute_score': 0.0  # 临时值
                            }
                    except (IndexError, ValueError, AttributeError):
                        continue

            # 检查是否有成绩数据
            if not data['scores']:
                raise ValueError("未找到有效成绩数据")

            # 计算课程目标平均分
            avg_score = np.mean(data['scores'])
            for obj in objectives.values():
                obj['absolute_score'] = round(avg_score * (obj['full_score'] / 100), 2)
                obj['evaluation'] = round(obj['absolute_score'] / obj['full_score'], 3)
                # 添加默认分析内容
                obj['analysis'] = f"课程目标{obj['index']}达成度为{obj['evaluation']}，详细分析待补充。"

            data['course_objectives'] = list(objectives.values())

            # 添加默认问题和改进措施
            data['problems'] = [
                "学生基础知识掌握不牢固",
                "综合分析能力有待提高",
                "解题思路不够清晰",
                "答题规范性不足"
            ]
            data['improvement_measures'] = [
                "加强基础知识的讲解和练习",
                "增加案例分析和讨论环节",
                "提供更多解题思路训练",
                "规范答题格式要求"
            ]

            return {
                'academic_year': data['学年'],
                'semester': data['学期'],
                'course_name': data['课程名称'],
                'teacher_name': data['命题人'],
                'class_name': data['考试班级'],
                'exam_date': data['考试日期'],
                'scores': data['scores'],
                'course_objectives': data['course_objectives'],
                'problems': data['problems'],
                'improvement_measures': data['improvement_measures']
            }

        except Exception as e:
            logging.error(f"Excel文件解析失败: {str(e)}", exc_info=True)
            return None

    def generate_score_distribution_chart(self, scores: List[int], filename: str) -> Optional[str]:
        """生成成绩分布图并保存"""
        try:
            plt.rcParams['font.sans-serif'] = ['SimHei']
            plt.rcParams['axes.unicode_minus'] = False

            fig, ax = plt.subplots(figsize=(8, 4))
            bins = [0, 60, 70, 80, 90, 100]
            labels = ['不及格', '60-69', '70-79', '80-89', '90-100']

            hist, _ = np.histogram(scores, bins=bins)
            bars = ax.bar(labels, hist, color='#4472C4')

            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width() / 2., height,
                        f'{int(height)}', ha='center', va='bottom')

            ax.set_title('成绩分布图', fontsize=12)
            ax.set_xlabel('分数段', fontsize=10)
            ax.set_ylabel('人数', fontsize=10)

            chart_path = os.path.join(self.output_dir, filename)
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()

            return chart_path
        except Exception as e:
            logging.error(f"生成成绩分布图失败: {str(e)}")
            return None

    def calculate_statistics(self, scores: List[int]) -> Dict:
        """计算成绩统计信息"""
        scores_array = np.array(scores)
        return {
            'average': round(np.mean(scores_array), 2),
            'excellent_rate': round(np.sum(scores_array >= 90) / len(scores) * 100, 2),
            'good_rate': round(np.sum((scores_array >= 80) & (scores_array < 90)) / len(scores) * 100, 2),
            'fail_rate': round(np.sum(scores_array < 60) / len(scores) * 100, 2)
        }

    def call_llm_for_analysis(self, prompt: str) -> Optional[str]:
        """调用大模型API获取分析结果"""
        try:
            response = Application.call(
                api_key=self.api_key,
                app_id=self.app_id,
                prompt=prompt
            )
            if response.status_code == 200:
                return response.output.text
            logging.error(f"API调用失败: {response.message}")
            return None
        except Exception as e:
            logging.error(f"API调用异常: {str(e)}")
            return None

    def generate_analysis_content(self, exam_data: Dict) -> Optional[Dict]:
        """生成试卷分析内容（调用大模型）"""
        # 准备提示词
        prompt = f"""请根据以下考试数据生成专业的试卷分析报告：

课程名称：{exam_data['course_name']}
考试班级：{exam_data['class_name']}
平均成绩：{self.calculate_statistics(exam_data['scores'])['average']}分

课程目标达成情况：
{'/n'.join([f"课程目标{obj['index']}：满分{obj['full_score']}分，平均得分{obj['absolute_score']}分" for obj in exam_data['course_objectives']])}

请按照以下结构生成分析报告：
1. 对每个课程目标的达成度分析和教学建议
2. 学生答题中存在的主要问题（至少列出4个）
3. 改进教学的具体措施（至少列出4条）

要求：
- 使用专业的教育评估术语
- 每个课程目标分析不少于100字
- 问题分析要具体
- 改进措施要具有可操作性"""

        # 调用大模型
        analysis = self.call_llm_for_analysis(prompt)
        if not analysis:
            logging.warning("大模型返回为空，使用默认分析内容")
            return exam_data  # 返回带有默认值的原始数据

        try:
            # 1. 解析课程目标分析部分
            objectives_section = re.search(r'1\..*?(课程目标[\s\S]*?)(?=2\.|\Z)', analysis, re.IGNORECASE)
            if objectives_section:
                objectives_text = objectives_section.group(1)
                for obj in exam_data['course_objectives']:
                    # 查找特定课程目标的分析
                    obj_pattern = re.compile(
                        fr'课程目标{obj["index"]}[：:](.*?)(?=课程目标|\Z)',
                        re.DOTALL
                    )
                    match = obj_pattern.search(objectives_text)
                    if match:
                        obj['analysis'] = match.group(1).strip()

            # 2. 解析问题分析部分
            problems_section = re.search(r'2\..*?主要问题([\s\S]*?)(?=3\.|\Z)', analysis, re.IGNORECASE)
            if problems_section:
                problems_text = problems_section.group(1)
                problems = [p.strip() for p in re.findall(r'[•\-]\s*(.*?)(?=\n|$)', problems_text) if p.strip()]
                if problems:
                    exam_data['problems'] = problems[:4] if len(problems) >= 4 else problems + [""] * (
                                4 - len(problems))

            # 3. 解析改进措施部分
            measures_section = re.search(r'3\..*?改进措施([\s\S]*?)(?=\Z)', analysis, re.IGNORECASE)
            if measures_section:
                measures_text = measures_section.group(1)
                measures = [m.strip() for m in re.findall(r'[•\-]\s*(.*?)(?=\n|$)', measures_text) if m.strip()]
                if measures:
                    exam_data['improvement_measures'] = measures[:4] if len(measures) >= 4 else measures + [""] * (
                                4 - len(measures))

            return exam_data
        except Exception as e:
            logging.error(f"解析大模型返回内容失败，使用默认分析: {str(e)}")
            return exam_data  # 即使解析失败也返回带有默认值的数据

    def generate_word_report(self, exam_data: Dict) -> Optional[str]:
        """生成Word格式的试卷分析表"""
        try:
            # 验证必要字段
            required_fields = ['academic_year', 'semester', 'course_name', 'teacher_name',
                               'class_name', 'exam_date', 'scores', 'course_objectives',
                               'problems', 'improvement_measures']
            for field in required_fields:
                if field not in exam_data:
                    raise ValueError(f"缺少必要字段: {field}")

            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)

            # 1. 标题
            title = doc.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.add_run("XX大学信息工程学院（人工智能学院）\n\n试 卷 分 析 表\n")
            title_run.font.size = Pt(16)
            title_run.bold = True

            semester = doc.add_paragraph()
            semester.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            semester.add_run(f"（{exam_data['academic_year']}学年第 {exam_data['semester']} 学期）\n")

            # 2. 基本信息表格
            basic_info = doc.add_table(rows=2, cols=4)
            basic_info.style = 'Table Grid'
            for row in basic_info.rows:
                for cell in row.cells:
                    cell.width = Inches(2.2)

            cells = basic_info.rows[0].cells
            cells[0].text = "课程名称"
            cells[1].text = exam_data['course_name']
            cells[2].text = "命题人"
            cells[3].text = exam_data['teacher_name']

            cells = basic_info.rows[1].cells
            cells[0].text = "考试班级"
            cells[1].text = exam_data['class_name']
            cells[2].text = "考试日期"
            cells[3].text = exam_data['exam_date']

            # 3. 成绩分析
            doc.add_paragraph("\n1、试卷成绩分析", style='Heading 1')
            doc.add_paragraph("（1）试卷成绩分布", style='Heading 2')

            stats = self.calculate_statistics(exam_data['scores'])
            excellent_good_rate = stats['excellent_rate'] + stats['good_rate']
            doc.add_paragraph(
                f"成绩分布较合理，平均成绩{stats['average']}分，"
                f"优良率{excellent_good_rate}%，不及格率{stats['fail_rate']}%。"
            )

            chart_path = self.generate_score_distribution_chart(exam_data['scores'], "score_distribution.png")
            if chart_path:
                doc.add_paragraph("成绩分布图如下：")
                doc.add_picture(chart_path, width=Inches(4.0))

            # 4. 课程目标分析
            doc.add_paragraph("\n2、试卷对课程目标的达成度分析", style='Heading 1')
            for obj in exam_data['course_objectives']:
                analysis_text = obj.get('analysis',
                                        f"课程目标{obj['index']}达成度为{obj['evaluation']}，详细分析待补充。")
                doc.add_paragraph(
                    f"（{obj['index']}）课程目标 {obj['index']}（支撑毕业要求 {obj['requirement']}）"
                    f"达成度为 {obj['evaluation']}，{analysis_text}"
                )

            # 5. 问题分析
            doc.add_paragraph("\n3、学生答题主要问题分析：", style='Heading 1')
            for i, problem in enumerate(exam_data.get('problems', []), 1):
                doc.add_paragraph(f"（{i}）{problem}")

            # 6. 改进措施
            doc.add_paragraph("\n4、存在问题及改进措施", style='Heading 1')
            for i, measure in enumerate(exam_data.get('improvement_measures', []), 1):
                doc.add_paragraph(f"（{i}）{measure}")

            # 签名
            doc.add_paragraph("\n\n签名：\t\t\t\t年\t月\t日")

            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(self.output_dir, f"{exam_data['course_name']}_{timestamp}_试卷分析表.docx")

            doc.save(output_path)
            return output_path
        except Exception as e:
            logging.error(f"生成Word报告失败: {str(e)}", exc_info=True)
            return None

    def process_exam_file(self, file_path: str) -> Optional[str]:
        """处理Excel文件并生成试卷分析表"""
        try:
            # 1. 读取Excel数据
            exam_data = self.read_excel_data(file_path)
            if not exam_data:
                logging.error("无法读取Excel数据")
                return None

            # 2. 调用大模型生成分析内容
            analyzed_data = self.generate_analysis_content(exam_data)
            if not analyzed_data:
                logging.error("无法生成分析内容")
                return None

            # 3. 生成Word报告
            result = self.generate_word_report(analyzed_data)
            if not result:
                logging.error("无法生成Word报告")
                return None

            return result
        except Exception as e:
            logging.error(f"处理考试文件失败: {str(e)}", exc_info=True)
            return None


# 示例使用
if __name__ == "__main__":
    # 替换为您的Excel文件路径
    excel_file = "exam_scores.xlsx"

    generator = ExamAnalysisGenerator()
    result = generator.process_exam_file(excel_file)

    if result:
        print(f"试卷分析表已成功生成，保存路径: {result}")
    else:
        print("试卷分析表生成失败，请查看日志文件获取详细信息。")