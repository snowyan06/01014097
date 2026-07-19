# -*- coding: utf-8 -*-
"""生成 EduGenius 教育平台技术文档 (Word格式)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ========== 全局样式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 51, 102)
    if level == 1:
        hs.font.size = Pt(22)
    elif level == 2:
        hs.font.size = Pt(16)
    else:
        hs.font.size = Pt(14)

def add_table(headers, rows):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    doc.add_paragraph()
    return table

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p

# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('EduGenius 智能教育平台')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('技 术 文 档')
run2.font.size = Pt(28)
run2.font.color.rgb = RGBColor(0, 102, 153)
run2.font.name = '黑体'
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = ver.add_run(f'版本：V1.0\n日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run3.font.size = Pt(14)

doc.add_page_break()

# ==================== 目录页 ====================
doc.add_heading('目  录', level=1)
toc_items = [
    '1. 项目概述',
    '2. 系统架构设计',
    '   2.1 整体架构',
    '   2.2 技术栈总览',
    '   2.3 模块划分',
    '3. 后端服务（Java / Spring Boot）',
    '   3.1 项目结构',
    '   3.2 实体类设计',
    '   3.3 RESTful API 接口',
    '4. AI 智能服务（Python / FastAPI）',
    '   4.1 项目结构',
    '   4.2 Agent 智能体架构',
    '   4.3 API 接口列表',
    '5. 前端应用（Vue 3）',
    '   5.1 项目结构',
    '   5.2 路由与页面',
    '   5.3 核心依赖',
    '6. 数字人对话与唇形同步技术',
    '   6.1 系统架构总览',
    '   6.2 核心技术栈',
    '   6.3 完整数据流',
    '   6.4 关键模块详解',
    '   6.5 API 接口汇总',
    '   6.6 数字人头像数据准备',
    '   6.7 性能指标与优化',
    '   6.8 核心文件索引',
    '7. 数据库设计',
    '   7.1 数据库表结构',
    '   7.2 表关系说明',
    '8. 环境配置与部署',
    '   8.1 环境变量配置',
    '   8.2 服务启动说明',
    '9. 总结',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.8

doc.add_page_break()

# ==================== 1. 项目概述 ====================
doc.add_heading('1. 项目概述', level=1)
doc.add_paragraph(
    'EduGenius 智能教育平台是一个面向高校教学场景的 AI 赋能教育系统，'
    '采用前后端分离 + 多微服务协同的架构设计。平台集成了大语言模型（LLM）能力，'
    '围绕"教、学、练、测、评"五大环节，为教师和学生提供智能化的教学辅助工具。'
)
doc.add_paragraph('平台核心功能包括：')
features = [
    '教师端：AI 辅助备课（自动生成教学计划）、智能出题、试卷自动处理（生成标准答案与评分标准）、试卷分析报表生成、教学效率分析',
    '学生端：AI 智能问答、在线练习与自动批改、知识图谱可视化、学习分析仪表盘、AI 模拟面试',
    '管理端：用户管理、教学资源管理、数据统计分析',
    '数字人交互：基于 WebRTC 的实时数字人对话，支持语音合成与唇形同步',
]
for f in features:
    add_bullet(f)

doc.add_paragraph()

# ==================== 2. 系统架构设计 ====================
doc.add_heading('2. 系统架构设计', level=1)
doc.add_heading('2.1 整体架构', level=2)
doc.add_paragraph(
    '系统采用多服务协同架构，由以下四个独立服务组成，通过 HTTP/REST API 进行通信：'
)
arch_rows = [
    ['EduGenius 前端', 'Vue 3 + Vite', '5173', '用户交互界面，SPA单页应用'],
    ['Edu_platform 后端', 'Spring Boot 3.2.5', '8080', '核心业务API，数据持久化，用户认证'],
    ['Edu_py AI服务', 'FastAPI + Uvicorn', '8000', 'AI智能体服务，LLM调用，文件处理'],
    ['avatar_py 数字人', 'Flask + aiohttp', '8010', '数字人实时交互，WebRTC推流'],
]
add_table(['服务模块', '技术框架', '默认端口', '职责说明'], arch_rows)

doc.add_heading('2.2 技术栈总览', level=2)
doc.add_heading('后端（Java）', level=3)
java_tech = [
    ['Spring Boot', '3.2.5', '核心应用框架'],
    ['Spring Data JPA', '3.2.5', 'ORM数据持久层'],
    ['MySQL Connector', '8.0.33', '数据库驱动'],
    ['Lombok', '1.18.32', '代码简化工具'],
    ['JDK', '17', 'Java运行环境'],
]
add_table(['技术', '版本', '用途'], java_tech)

doc.add_heading('AI服务（Python）', level=3)
py_tech = [
    ['FastAPI', '0.116.1', '异步Web框架'],
    ['Uvicorn', '0.33.0', 'ASGI服务器'],
    ['DashScope SDK', '>=1.20.0', '通义千问大模型调用'],
    ['OpenAI SDK', '>=1.50.0', '兼容OpenAI协议的模型调用'],
    ['PyMySQL', '>=1.0.2', 'MySQL数据库连接'],
    ['python-docx', '>=1.1.0', 'Word文档生成'],
    ['openpyxl', '>=3.0.7', 'Excel文件处理'],
    ['PyPDF2', '>=3.0.1', 'PDF文件解析'],
    ['Transformers', '4.46.2', '本地大模型推理（可选）'],
]
add_table(['技术', '版本', '用途'], py_tech)

doc.add_heading('前端', level=3)
fe_tech = [
    ['Vue', '3.4.34', '前端框架'],
    ['Vite', '5.4.11', '构建工具'],
    ['PrimeVue', '4.0.0', 'UI组件库'],
    ['TailwindCSS', '3.4.6', '原子化CSS框架'],
    ['Axios', '1.7.7', 'HTTP客户端'],
    ['ECharts', '5.6.0', '数据可视化图表'],
    ['Chart.js', '3.3.2', '图表库'],
    ['D3.js', '7.9.0', '知识图谱可视化'],
    ['vis-network', '10.0.1', '网络图/关系图'],
    ['Pinia', '3.0.4', '状态管理'],
    ['Vue Router', '4.4.0', '路由管理'],
]
add_table(['技术', '版本', '用途'], fe_tech)

doc.add_heading('数字人系统', level=3)
av_tech = [
    ['PyTorch', 'latest', '深度学习框架'],
    ['MuseTalk', '-', '唇形同步生成模型'],
    ['Edge TTS', '-', '微软语音合成'],
    ['aiortc', '-', 'WebRTC实现'],
    ['Flask + aiohttp', '-', 'HTTP/WebSocket服务'],
]
add_table(['技术', '版本', '用途'], av_tech)

doc.add_heading('2.3 模块划分', level=2)
doc.add_paragraph('项目目录结构如下：')
structure = [
    '01014097/',
    '├── EduGenius/          # Vue 3 前端应用',
    '├── Edu_platform/       # Spring Boot Java 后端',
    '├── Edu_py/             # Python AI 智能服务',
    '├── avatar_py/          # 数字人交互系统',
    '├── backend.py          # Flask 认证服务（备用）',
    '├── mydb.sql            # 数据库初始化脚本',
    '├── .env                # 全局环境变量配置',
    '└── quick_start.ps1     # 一键启动脚本',
]
for line in structure:
    p = doc.add_paragraph(line)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    for r in p.runs:
        r.font.name = 'Consolas'
        r.font.size = Pt(10)

doc.add_page_break()

# ==================== 3. 后端服务 ====================
doc.add_heading('3. 后端服务（Java / Spring Boot）', level=1)
doc.add_heading('3.1 项目结构', level=2)
doc.add_paragraph('Java 后端采用经典的三层架构（Controller - Service - Repository），包路径为 org.example：')
java_struct = [
    ['config/', '配置类', 'CorsConfig（跨域）、DataInitializer（数据初始化）、FileStorageConfig（文件存储）'],
    ['controller/', '控制层（10个）', 'AuthController、QuestionController、UserAnswerController、PracticeSessionController、TeachingMaterialController、TeachingEfficiencyController、UserAdminController、UserProfileController、FileDownloadController'],
    ['service/', '业务层', 'AuthService、QuestionService、UserAnswerService 等，封装核心业务逻辑'],
    ['repository/', '数据访问层', '基于 Spring Data JPA 的 Repository 接口'],
    ['entity/', '实体类', 'User、Question、UserAnswer、PracticeSession、TeachingMaterial、TeachingEfficiency、KnowledgePoint 等'],
    ['dto/', '数据传输对象', 'LoginRequest/Response、UserAnswerStatsDto、CorrectRateTrendDto、KnowledgeMasteryDto、FrequentErrorDto、TeacherActivityDTO 等'],
    ['exception/', '异常处理', '自定义异常类与全局异常处理器'],
]
add_table(['包名', '说明', '内容'], java_struct)

doc.add_heading('3.2 实体类设计', level=2)
doc.add_paragraph('系统核心实体类与数据库表一一对应，使用 JPA 注解进行 ORM 映射。主要实体包括：')
entities = [
    ['User', 'users', '用户信息（用户名、密码、角色、昵称）'],
    ['Question', 'questions', '题库（题目ID、类型、内容、答案、难度、来源）'],
    ['UserAnswer', 'user_answers', '学生答题记录（题目、作答、批改结果、会话ID）'],
    ['PracticeSession', 'practice_sessions', '练习会话（用户、标题、正确率、分数）'],
    ['TeachingMaterial', 'teaching_materials', '教学资源（文件名、类型、大小、存储路径）'],
    ['TeachingEfficiency', 'teaching_efficiency', '教学效率（备课时间、修改次数、优化建议）'],
]
add_table(['实体类', '对应表', '说明'], entities)

doc.add_heading('3.3 RESTful API 接口', level=2)
doc.add_paragraph('Java 后端提供以下 RESTful API 接口（基础路径：http://localhost:8080）：')

doc.add_heading('认证模块 /api/auth', level=3)
add_table(['方法', '路径', '说明'], [
    ['POST', '/api/auth/login', '用户登录（支持 student/teacher/admin 三种角色）'],
])

doc.add_heading('题目管理 /api/questions', level=3)
add_table(['方法', '路径', '说明'], [
    ['POST', '/api/questions', '创建题目'],
    ['GET', '/api/questions', '获取所有题目'],
    ['GET', '/api/questions/{id}', '按ID获取题目'],
    ['GET', '/api/questions/by-number/{questionId}', '按题目编号获取'],
    ['PUT', '/api/questions/{id}', '更新题目'],
    ['DELETE', '/api/questions/{id}', '删除题目'],
    ['GET', '/api/questions/daily', '教师日出题统计'],
    ['GET', '/api/questions/weekly', '教师周出题统计'],
    ['GET', '/api/questions/total', '教师总出题统计'],
])

doc.add_heading('答题记录 /api/user-answers', level=3)
add_table(['方法', '路径', '说明'], [
    ['POST', '/api/user-answers', '创建答题记录'],
    ['GET', '/api/user-answers', '获取所有记录'],
    ['GET', '/api/user-answers/by-session/{sessionId}', '按会话查询'],
    ['GET', '/api/user-answers/user/{userId}', '按用户查询'],
    ['GET', '/api/user-answers/user-stats', '用户答题统计'],
    ['GET', '/api/user-answers/stats/correct-rate-trend/{userId}', '正确率趋势'],
    ['GET', '/api/user-answers/stats/knowledge-mastery/{userId}', '知识点掌握度'],
    ['GET', '/api/user-answers/stats/frequent-errors/{userId}', '高频错题'],
])

doc.add_heading('其他模块', level=3)
add_table(['方法', '路径', '说明'], [
    ['GET/POST', '/api/practice-sessions/*', '练习会话管理（CRUD）'],
    ['GET/POST', '/api/teaching-materials/*', '教学资源管理（含文件上传）'],
    ['GET/POST', '/api/teaching-efficiency/*', '教学效率分析'],
    ['GET/POST', '/api/admin/users/*', '用户管理（管理员）'],
    ['GET', '/api/user-profile/{userId}', '用户资料'],
    ['GET', '/api/files/{filename}', '文件下载'],
])

doc.add_page_break()

# ==================== 4. AI 智能服务 ====================
doc.add_heading('4. AI 智能服务（Python / FastAPI）', level=1)
doc.add_heading('4.1 项目结构', level=2)
py_struct = [
    ['agents/', 'AI智能体', 'TeachingPlanAgent、QAAgent、QuestionGenAgent、StudentPracticeAgent、ExamProcessAgent、ExamAnalysisAgent、KnowledgeGraphAgent'],
    ['adapters/', '模型适配器', 'LLMAdapter（统一LLM调用接口）、IFlytekAdapter（讯飞适配器）'],
    ['config/', '配置管理', 'Settings（全局配置）、logging_config（日志配置）'],
    ['services/', '基础服务', 'STTService（语音识别）、TTSService（语音合成）、output_filter（输出过滤）'],
    ['rag/', 'RAG检索', 'indexer（文档索引）、retriever（知识检索）'],
    ['main.py', 'API入口', 'FastAPI应用主入口，注册所有API路由'],
    ['student_app.py', '学生端模块', 'AI出题与自动批改（DashScope Agent调用）'],
    ['teacher_app.py', '教师端模块', '教师端AI出题（DashScope Agent调用）'],
    ['document.py', '文档处理', 'DocumentProcessor（docx/xlsx/txt解析）'],
]
add_table(['文件/目录', '说明', '内容'], py_struct)

doc.add_heading('4.2 Agent 智能体架构', level=2)
doc.add_paragraph(
    'AI 服务采用 Agent（智能体）架构模式，每个 Agent 继承自 BaseAgent 基类，'
    '通过统一的 LLM Adapter 调用大语言模型。系统支持多模型切换（通义千问 DashScope / 讯飞星火），'
    '通过环境变量 LLM_PROVIDER 配置。'
)
agents_info = [
    ['TeachingPlanAgent', '教学计划生成', '根据课程大纲自动生成结构化教学计划'],
    ['QAAgent', '智能问答', '基于上下文的教学问答，支持RAG检索增强'],
    ['QuestionGenAgent', '题目生成（教师）', '根据提示词自动生成各类型题目'],
    ['StudentPracticeAgent', '学生练习与批改', '生成题目 + 自动批改 + 反馈解析'],
    ['ExamProcessAgent', '试卷处理', '解析试卷文件，自动生成标准答案和评分标准'],
    ['ExamAnalysisAgent', '试卷分析', '分析考试成绩Excel，生成试卷分析报告（docx）'],
    ['KnowledgeGraphAgent', '知识图谱', '基于答题记录构建个人知识图谱，计算掌握度'],
]
add_table(['Agent 名称', '功能定位', '说明'], agents_info)

doc.add_heading('4.3 API 接口列表', level=2)
doc.add_paragraph('AI 服务运行在 http://localhost:8000，提供以下接口：')
py_apis = [
    ['POST', '/generate-course', '生成教学计划', 'CourseOutlineRequest'],
    ['POST', '/ask-question', '智能问答', 'QuestionRequest'],
    ['POST', '/generate-question', '教师端出题', 'QuestionGenerateRequest'],
    ['POST', '/question/generate', '学生端出题', 'QuestionRequest_stu'],
    ['POST', '/question/grade', '自动批改', 'AnswerGradeRequest'],
    ['POST', '/upload-file', '文件上传解析', 'multipart/form-data'],
    ['POST', '/process-exam', '试卷处理', 'multipart/form-data (pdf/docx/txt)'],
    ['POST', '/analyze-exam', '试卷分析', 'multipart/form-data (xlsx)'],
    ['GET', '/download-file/{filename}', '文件下载', '-'],
    ['GET', '/knowledge-graph', '获取知识图谱', 'user_id'],
    ['GET', '/user-knowledge-mastery', '知识点掌握度', 'user_id'],
    ['POST', '/generate-mindmap', '生成思维导图', 'MultimodalRequest'],
    ['POST', '/generate-code-examples', '生成代码案例', 'MultimodalRequest'],
    ['POST', '/generate-practice-questions', '生成练习题', 'MultimodalRequest'],
    ['POST', '/generate-analysis-report', '生成解析报告', 'MultimodalRequest'],
]
add_table(['方法', '路径', '说明', '请求参数'], py_apis)

doc.add_page_break()

# ==================== 5. 前端应用 ====================
doc.add_heading('5. 前端应用（Vue 3）', level=1)
doc.add_heading('5.1 项目结构', level=2)
fe_struct = [
    ['src/views/pages/', '页面视图', 'Landing（首页）、TeacherCenter、StudentAssistant、AIMockInterview、AnalysisImprovement、AnalyticsDashboard 等'],
    ['src/views/pages/auth/', '认证页面', 'Login（登录）、Register（注册）、Access（权限）、Error（错误）'],
    ['src/components/', '公共组件', 'DashboardOverview、DigitalHuman、LearningAnalyticsDashboard、DetailModal、PracticeHistoryDetailModal、SelectionAssistant、LoginBackground 等'],
    ['src/layout/', '布局组件', 'AppLayout（主布局）、侧边栏、顶栏等'],
    ['src/router/', '路由配置', 'Vue Router 路由定义'],
    ['src/store/', '状态管理', 'Pinia 全局状态'],
    ['src/assets/', '静态资源', '样式、图片、布局资源'],
]
add_table(['目录', '说明', '内容'], fe_struct)

doc.add_heading('5.2 路由与页面', level=2)
routes = [
    ['/', 'dashboard', '首页/仪表盘', 'Landing.vue'],
    ['/TeacherCenter', 'TeacherCenter', '教师中心', 'TeacherCenter.vue'],
    ['/StudentAssistant', 'StudentAssistant', '学生助手', 'StudentAssistant.vue'],
    ['/AIMockInterview', 'AIMockInterview', 'AI模拟面试', 'AIMockInterview.vue'],
    ['/AnalysisImprovement', 'AnalysisImprovement', '分析改进', 'AnalysisImprovement.vue'],
    ['/Analytics', 'Analytics', '数据分析面板', 'AnalyticsDashboard.vue'],
    ['/documentation', 'documentation', '文档中心', 'Documentation.vue'],
    ['/account', 'account', '个人账户', 'Account.vue'],
    ['/settings', 'settings', '系统设置', 'Settings.vue'],
    ['/login', 'login', '登录页', 'Login.vue'],
    ['/register', 'register', '注册页', 'register.vue'],
]
add_table(['路径', '路由名', '页面说明', '组件文件'], routes)

doc.add_heading('5.3 核心依赖', level=2)
doc.add_paragraph('前端核心第三方依赖：')
fe_deps = [
    'PrimeVue 4.0 - 企业级 UI 组件库，提供表格、对话框、表单等丰富组件',
    'TailwindCSS 3.4 - 原子化 CSS 框架，实现快速样式开发',
    'ECharts 5.6 - 百度开源图表库，用于数据分析可视化',
    'D3.js 7.9 / vis-network 10.0 - 知识图谱与关系网络可视化',
    'Axios 1.7 - HTTP 请求库，与后端 API 通信',
    'Marked 16.1 - Markdown 渲染，用于 AI 回复内容展示',
    'mammoth 1.9 - Word 文档前端解析',
    'file-saver + JSZip - 文件下载与压缩包处理',
    'WebRTC Adapter - 数字人实时音视频通信',
]
for dep in fe_deps:
    add_bullet(dep)

doc.add_page_break()

# ==================== 6. 数字人对话与唇形同步技术 ====================
doc.add_heading('6. 数字人对话与唇形同步技术', level=1)

# 6.1 系统架构总览
doc.add_heading('6.1 系统架构总览', level=2)
doc.add_paragraph(
    '本系统实现了一套完整的实时数字人对话系统，支持语音输入、AI对话、语音合成（TTS）、'
    '唇形驱动渲染的全链路闭环。系统分为三层架构：'
)
arch_layers = [
    '前端层（Vue3）：WebRTC 视频/音频接收、语音录制与 STT 上传、对话 UI 交互、绿幕抠除（可选）',
    '信令服务层（Python aiohttp）：WebRTC 信令交换、文字转 TTS + 唇形驱动、AI 对话自动播报、语音识别、说话状态轮询',
    '数字人渲染层（Python）：TTS 引擎、ASR 音频特征提取（Whisper）、唇形推理（MuseTalk UNet + VAE）、帧融合、WebRTC 推送',
]
for layer in arch_layers:
    add_bullet(layer)

# 6.2 核心技术栈
doc.add_heading('6.2 核心技术栈', level=2)
av_core_tech = [
    ['唇形驱动模型', 'MuseTalk', '基于 UNet + VAE 的音频驱动唇形生成'],
    ['音频特征提取', 'Whisper (tiny)', '将音频转为 Mel 频谱特征，供 UNet 条件推理'],
    ['语音合成 (TTS)', 'EdgeTTS (默认)', '微软 Edge TTS，支持多种中文音色'],
    ['语音识别 (STT)', 'faster-whisper / openai-whisper', '本地 Whisper 模型，支持中英文'],
    ['AI 对话', '通义千问 (DashScope)', 'qwen-turbo 模型，支持多轮对话'],
    ['实时传输', 'WebRTC (aiortc)', '低延迟音视频传输'],
    ['信令服务器', 'aiohttp', 'Python 异步 HTTP/WebSocket 服务'],
    ['前端框架', 'Vue3 + Pinia', '响应式 UI + 全局状态管理'],
    ['面部分割', 'FaceParsing (BiSeNet)', '生成唇部区域 mask，用于帧融合'],
    ['视频编码', 'OpenCV + PyAV', '帧处理与 WebRTC VideoFrame 转换'],
]
add_table(['模块', '技术选型', '说明'], av_core_tech)

# 6.3 完整数据流
doc.add_heading('6.3 完整数据流', level=2)

doc.add_heading('6.3.1 对话流程（用户说话 → 数字人回复）', level=3)
doc.add_paragraph('完整对话数据流如下：')
conv_flow = [
    '1. 用户语音 → 前端录音 → WebM/WAV 音频',
    '2. POST /speech_to_text → WhisperSTTService.transcribe_audio() → 返回识别文本',
    '3. 前端显示用户文本',
    '4. POST /chat {sessionid, message} → AIChatService.get_ai_response() → 调用通义千问 API',
    '5. nerfreal.put_msg_txt(ai_response) → TTS引擎合成 → 唇形推理 → 帧融合 → WebRTC推送',
    '6. 前端 <video> 元素播放 → 用户看到数字人说话+听到声音',
]
for step in conv_flow:
    add_bullet(step)

doc.add_heading('6.3.2 直接播报流程（文字 → 数字人说话）', level=3)
doc.add_paragraph('直接播报数据流如下：')
broadcast_flow = [
    '1. POST /human {sessionid, text, interrupt}',
    '2. nerfreal.put_msg_txt(text) → EdgeTTS.txt_to_audio() 流式合成 MP3',
    '3. pydub 解码为 PCM (16kHz, mono) → 按 chunk=320 采样点（20ms）切片',
    '4. BaseReal.put_audio_frame() → 发送到 ASR 队列 + WebRTC 音频轨道',
    '5. MuseASR 提取 Whisper 特征 → MuseTalk Inference 生成唇形帧',
    '6. process_frames 融合帧 → WebRTC VideoTrack 推送',
]
for step in broadcast_flow:
    add_bullet(step)

# 6.4 关键模块详解
doc.add_heading('6.4 关键模块详解', level=2)

doc.add_heading('6.4.1 WebRTC 信令建立连接', level=3)
doc.add_paragraph(
    '前端创建 RTCPeerConnection 并生成 Offer SDP，通过 POST /offer 发送到服务端。'
    '服务端创建数字人实例（MuseReal）和 HumanPlayer（音频+视频轨道），'
    '完成 SDP 交换后建立 ICE 连接，形成 WebRTC 双向媒体通道。'
    '每个 session 独立维护 RTCPeerConnection 和数字人实例。'
)

doc.add_heading('6.4.2 TTS 语音合成', level=3)
doc.add_paragraph(
    'EdgeTTS 模块通过 edge_tts.Communicate 流式获取 MP3 数据，使用 pydub 解码为 PCM '
    '（16kHz, 单声道, 16-bit），按 320 采样点（20ms）切片后逐帧发送。'
)
doc.add_paragraph('支持的 TTS 引擎：')
tts_engines = [
    ['EdgeTTS', '云端', 'edge-tts Python库', '默认，免费，多音色'],
    ['GPT-SoVITS', '32kHz→16kHz', 'HTTP POST', '开源克隆 TTS'],
    ['CosyVoice', '24kHz→16kHz', 'HTTP GET', '阿里开源'],
    ['FishTTS', '44.1kHz→16kHz', 'HTTP POST', '开源 TTS'],
    ['TencentTTS', '16kHz', 'HTTP POST', '腾讯云 TTS'],
    ['DoubaoTTS', '16kHz', 'WebSocket', '火山引擎 TTS'],
    ['XTTS', '24kHz→16kHz', 'HTTP POST', 'Coqui 开源'],
]
add_table(['引擎', '采样率', '接口方式', '说明'], tts_engines)

doc.add_heading('6.4.3 音频特征提取 (ASR)', level=3)
doc.add_paragraph(
    'MuseASR 模块将 PCM 音频帧转换为 Whisper 音频特征向量。'
    '工作流程：从音频帧队列累积到 batch_size 帧后批量处理，通过 Whisper encoder 提取 '
    'Mel 频谱特征，按帧率切分为每帧对应的特征片段。输出形状为 (batch, 50, 384) 的特征张量。'
)
doc.add_paragraph('关键参数：音频采样率 16000Hz，Whisper 模型 tiny (39M参数)，特征维度 384，帧率 15fps。')

doc.add_heading('6.4.4 唇形推理 (MuseTalk Inference)', level=3)
doc.add_paragraph(
    '核心模型包括 VAE（Stable Diffusion VAE，图像编解码）、UNet（音频条件唇形生成）'
    '和 PositionalEncoding（时间位置编码）。'
)
doc.add_paragraph('推理流程：')
inference_steps = [
    '1. whisper_chunks（音频特征）→ PositionalEncoding 添加时间位置信息',
    '2. 从 input_latent_list_cycle 取出对应帧的 VAE 编码 latent',
    '3. UNet 推理：输入 latent_batch + audio_feature_batch + timesteps → 输出 pred_latents',
    '4. VAE 解码：pred_latents → res_frames（256x256 RGB 图像）',
    '5. 放入 res_frame_queue 等待融合',
]
for step in inference_steps:
    add_bullet(step)
doc.add_paragraph('8GB 显存优化：batch_size 强制不超过 2，模型使用 FP16 半精度，推理完成后立即释放显存。')

doc.add_heading('6.4.5 帧融合 (paste_back_frame)', level=3)
doc.add_paragraph(
    '帧融合模块将 MuseTalk 生成的唇形区域无缝融合到原始帧上。'
    '融合流程：RGB→BGR 颜色空间转换 → resize 到 bbox 大小 → 读取预生成的 RGBA mask '
    '（FaceParsing 生成的面部区域 mask）→ 用 alpha 通道控制混合（内部显示唇形输出，'
    '外部显示原始帧，边缘高斯模糊过渡 + 肤色渐变）→ 贴回原图坐标位置。'
)
doc.add_paragraph(
    'mask 生成（genavatar_musetalk.py）：使用 FaceParsing (BiSeNet) 分割面部区域，'
    '从原图提取真实肤色填充 mask，使用距离变换创建渐变过渡边缘，保存为 RGBA 格式 PNG。'
)

doc.add_heading('6.4.6 WebRTC 帧推送', level=3)
doc.add_paragraph(
    '渲染线程将融合后的帧通过 WebRTC 推送。视频帧通过 VideoFrame.from_ndarray() 转换，'
    '音频帧通过 AudioFrame 构造（s16 格式，mono，16kHz）。'
    'WebRTC 轨道实现：HumanVideoTrack（异步队列驱动）和 HumanAudioTrack（精确时间戳管理，90kHz 时钟）。'
)

doc.add_heading('6.4.7 前端接收与播放', level=3)
doc.add_paragraph(
    '前端通过 RTCPeerConnection 建立 WebRTC 连接，监听远程媒体流并绑定到 <video> 元素。'
    '支持 TTS 触发（POST /human）、说话状态轮询（POST /is_speaking，500ms 间隔）。'
    '可选绿幕抠除：使用 WebGL 着色器实现实时色度键控，将绿色背景替换为透明。'
)

doc.add_heading('6.4.8 语音识别 (STT)', level=3)
doc.add_paragraph(
    '前端录音（MediaRecorder API）→ WebM 格式 → 转换为 WAV（单声道, 16kHz）→ '
    'POST /speech_to_text 上传 → 后端使用 faster-whisper / openai-whisper 识别 → 返回文本。'
    '优先使用 faster-whisper（CTranslate2, CPU 友好, int8 量化），回退到 openai-whisper。'
    '模型大小 tiny (39M)，启用 VAD 过滤（min_silence_duration_ms=500）。'
)

# 6.5 API 接口汇总
doc.add_heading('6.5 API 接口汇总', level=2)
doc.add_paragraph('数字人系统完整 API 接口列表（http://localhost:8010）：')
avatar_apis = [
    ['POST', '/offer', 'WebRTC 信令交换', '{sdp, type, sessionid}'],
    ['POST', '/human', '文字转 TTS 播报', '{sessionid, text, interrupt}'],
    ['POST', '/chat', 'AI 对话 + 自动播报', '{sessionid, message}'],
    ['POST', '/speech_to_text', '语音识别', 'FormData: audio, language'],
    ['POST', '/is_speaking', '查询说话状态', '{sessionid}'],
    ['POST', '/interrupt', '打断当前播报', '{sessionid}'],
    ['POST', '/interview_chat', '面试官对话', '{sessionid, message, config}'],
    ['POST', '/interview_question', '生成面试问题', '{sessionid, type, position, difficulty}'],
    ['POST', '/end_interview', '结束面试+评估', '{sessionid, config, signal_summary}'],
    ['POST', '/parse_resume', '解析简历文件', 'FormData: resume'],
    ['POST', '/extract_keywords', '提取关键词', '{text, position}'],
]
add_table(['方法', '路径', '说明', '请求参数'], avatar_apis)

# 6.6 数字人头像数据准备
doc.add_heading('6.6 数字人头像数据准备', level=2)
doc.add_paragraph('生成数字人头像需要以下预处理步骤（genavatar_musetalk.py）：')
avatar_prep = [
    '1. 视频帧提取：从输入视频按 fps 提取帧图片 → full_imgs/',
    '2. 人脸检测：YOLOv8 检测每帧人脸 bbox',
    '3. VAE 编码：将每帧人脸区域通过 VAE 编码为 latent → latents.pt',
    '4. 坐标记录：记录每帧人脸 bbox 坐标 → coords.pkl',
    '5. Mask 生成：FaceParsing 分割面部 → RGBA 肤色 mask → mask/',
    '6. Mask 坐标：记录 mask 裁剪坐标 → mask_coords.pkl',
]
for step in avatar_prep:
    add_bullet(step)
doc.add_paragraph('头像目录结构：data/avatars/{avatar_id}/，包含 full_imgs/、coords.pkl、latents.pt、mask/、mask_coords.pkl、vid_output/、avator_info.json。')

# 6.7 性能指标与优化
doc.add_heading('6.7 性能指标与优化', level=2)
doc.add_heading('延迟链路分析', level=3)
doc.add_paragraph(
    'TTS 合成(~200ms) → 首帧音频(~20ms) → ASR 特征提取(~10ms) → UNet 推理(~30ms) → '
    'VAE 解码(~15ms) → 帧融合(~5ms) → WebRTC 传输(~50ms) → 前端渲染(~16ms)。'
    '总延迟约 350-500ms（取决于网络和 GPU 性能）。'
)
doc.add_heading('性能埋点', level=3)
doc.add_paragraph(
    '前端 SelectionAssistant.vue 内置全链路性能埋点：'
    'ttsFetchStart（TTS请求发起）、ttsResponseTime（TTS响应返回）、'
    'firstFrameTime（WebRTC首帧到达）、firstRenderTime（首帧渲染完成）。'
)
doc.add_heading('显存优化（8GB GPU）', level=3)
gpu_opts = [
    'UNet + VAE 使用 FP16 半精度推理',
    'batch_size 限制最大为 2',
    '推理后立即 torch.cuda.empty_cache()',
    '模型预热 (warm_up) 避免首次推理延迟',
]
for opt in gpu_opts:
    add_bullet(opt)

# 6.8 核心文件索引
doc.add_heading('6.8 核心文件索引', level=2)
core_files = [
    ['avatar_py/webrtc_server_simple.py', 'WebRTC 信令服务器，所有 HTTP API 路由'],
    ['avatar_py/basereal.py', '数字人基类，TTS/音频帧/WebRTC 桥接'],
    ['avatar_py/musereal.py', 'MuseTalk 数字人实现，唇形推理+帧融合'],
    ['avatar_py/ttsreal.py', 'TTS 引擎集合 (EdgeTTS/GPT-SoVITS/...)'],
    ['avatar_py/museasr.py', 'Whisper 音频特征提取'],
    ['avatar_py/webrtc.py', 'WebRTC 音视频轨道实现'],
    ['avatar_py/ai_chat.py', '通义千问 AI 对话服务'],
    ['avatar_py/stt_service.py', 'Whisper 语音识别服务'],
    ['avatar_py/genavatar_musetalk.py', '头像数据预处理工具'],
    ['EduGenius/src/components/SelectionAssistant.vue', '前端数字人悬浮卡片组件'],
    ['EduGenius/src/components/DigitalHuman.vue', '前端数字人内嵌组件'],
    ['EduGenius/src/store/digitalHumanStore.js', 'Pinia 全局状态管理'],
]
add_table(['文件', '职责'], core_files)

doc.add_page_break()

# ==================== 7. 数据库设计 ====================
doc.add_heading('7. 数据库设计', level=1)
doc.add_paragraph('系统使用 MySQL 数据库（默认库名：mydb），字符集 utf8mb4。以下是核心数据表设计：')

doc.add_heading('7.1 数据库表结构', level=2)

doc.add_heading('users（用户表）', level=3)
add_table(['字段', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '用户ID'],
    ['username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '用户名'],
    ['nickname', 'VARCHAR(100)', '', '昵称'],
    ['password', 'VARCHAR(100)', 'NOT NULL', '密码'],
    ['role', 'VARCHAR(50)', 'NOT NULL', '角色（student/teacher/admin）'],
])

doc.add_heading('questions（题目表）', level=3)
add_table(['字段', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '自增主键'],
    ['question_id', 'VARCHAR(50)', 'NOT NULL', '题目编号'],
    ['type', 'VARCHAR(100)', 'NOT NULL', '题目类型（单选/判断/填空/简答）'],
    ['content', 'TEXT', 'NOT NULL', '题目内容'],
    ['answer', 'TEXT', 'NOT NULL', '标准答案'],
    ['difficulty', 'VARCHAR(50)', 'NOT NULL', '难度（简单/中等/困难）'],
    ['source', 'VARCHAR(100)', '', '来源（AI生成/手动添加）'],
    ['teacher_id', 'BIGINT', 'FK', '出题教师ID'],
    ['created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
])

doc.add_heading('user_answers（答题记录表）', level=3)
add_table(['字段', '类型', '约束', '说明'], [
    ['id', 'INT', 'PK, AUTO_INCREMENT', '自增主键'],
    ['question_type', 'VARCHAR(20)', 'NOT NULL', '题目类型'],
    ['question_content', 'TEXT', 'NOT NULL', '题目内容'],
    ['difficulty', 'VARCHAR(10)', 'NOT NULL', '难度等级'],
    ['user_answer', 'TEXT', '', '学生答案'],
    ['correct_answer', 'TEXT', '', '标准答案'],
    ['explanation', 'TEXT', '', '解析说明'],
    ['is_correct', 'TINYINT(1)', 'DEFAULT 0', '是否正确'],
    ['session_id', 'VARCHAR(36)', '', '练习会话UUID'],
    ['user_id', 'BIGINT', 'NOT NULL', '学生ID'],
])

doc.add_heading('practice_sessions（练习会话表）', level=3)
add_table(['字段', '类型', '约束', '说明'], [
    ['id', 'INT', 'PK, AUTO_INCREMENT', '自增主键'],
    ['user_id', 'INT', 'NOT NULL', '用户ID'],
    ['title', 'VARCHAR(100)', 'NOT NULL', '练习标题'],
    ['total_questions', 'INT', 'NOT NULL', '总题数'],
    ['correct_count', 'INT', 'NOT NULL', '正确数'],
    ['accuracy_rate', 'DECIMAL(5,2)', '', '正确率'],
    ['score', 'INT', '', '得分'],
    ['session_id', 'VARCHAR(36)', '', '会话UUID'],
    ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '创建时间'],
])

doc.add_heading('knowledge_points / knowledge_dependencies（知识图谱）', level=3)
add_table(['表名', '核心字段', '说明'], [
    ['knowledge_points', 'id, name, user_id', '知识点表，按用户隔离，name+user_id唯一'],
    ['knowledge_dependencies', 'id, parent, child, user_id', '知识点依赖关系（父子），外键关联knowledge_points'],
])

doc.add_heading('teaching_materials（教学资源表）', level=3)
add_table(['字段', '类型', '说明'], [
    ['id', 'BIGINT', '自增主键'],
    ['resource_name', 'VARCHAR(255)', '资源名称'],
    ['file_type', 'VARCHAR(50)', '文件类型（PDF/DOCX/TXT）'],
    ['teacher_id', 'BIGINT', '创建教师ID'],
    ['file_size', 'INT', '文件大小（字节）'],
    ['file_path', 'VARCHAR(512)', '存储路径'],
    ['created_at', 'DATETIME', '创建时间'],
])

doc.add_heading('teaching_efficiency（教学效率表）', level=3)
add_table(['字段', '类型', '说明'], [
    ['id', 'BIGINT', '自增主键'],
    ['teacher_id', 'BIGINT', '教师ID（外键关联users）'],
    ['date', 'DATE', '日期'],
    ['prep_time', 'INT', '备课时间（分钟）'],
    ['prep_revisions', 'INT', '备课修改次数'],
    ['optimization_notes', 'TEXT', '优化建议'],
])

doc.add_heading('7.2 表关系说明', level=2)
relations = [
    'users ←→ user_answers：一对多（一个用户多条答题记录）',
    'users ←→ practice_sessions：一对多',
    'users ←→ teaching_materials：一对多（教师创建教学资源）',
    'users ←→ teaching_efficiency：一对多（教师教学效率记录）',
    'users ←→ questions：一对多（教师出题）',
    'knowledge_points ←→ knowledge_dependencies：通过 (name, user_id) 外键关联，构建有向图',
    'user_answers.session_id ←→ practice_sessions.session_id：通过UUID关联练习详情',
]
for r in relations:
    add_bullet(r)

doc.add_page_break()

# ==================== 8. 环境配置与部署 ====================
doc.add_heading('8. 环境配置与部署', level=1)
doc.add_heading('8.1 环境变量配置', level=2)
doc.add_paragraph('项目通过根目录 .env 文件统一管理环境变量，关键配置如下：')
env_rows = [
    ['DB_HOST / DB_PORT', 'localhost / 3306', 'MySQL数据库地址'],
    ['DB_USER / DB_PASSWORD', 'root / 123456', '数据库账号密码'],
    ['DB_NAME', 'mydb', '数据库名称'],
    ['DASH_SCOPE_API_KEY', 'sk-xxx', '通义千问 DashScope API密钥'],
    ['OTHER_API_KEY', 'sk-xxx', '备用AI服务API密钥'],
    ['LLM_PROVIDER', 'dashscope', 'LLM提供商（dashscope/iflytek）'],
    ['LLM_MODEL', 'qwen-turbo', '使用的模型名称'],
    ['SERVER_PORT', '8080', 'Java后端端口'],
    ['AVATAR_PORT', '8010', '数字人服务端口'],
    ['TTS_SERVER', 'http://127.0.0.1:9880', 'TTS语音合成服务地址'],
    ['FILE_STORAGE_PATH', './storage/files', '文件存储路径'],
    ['VUE_APP_API_BASE_URL', 'http://localhost:8080', '前端→Java后端地址'],
    ['VUE_APP_PYTHON_API_URL', 'http://localhost:8000', '前端→AI服务地址'],
    ['VUE_APP_AVATAR_URL', 'http://localhost:8010', '前端→数字人地址'],
]
add_table(['变量名', '示例值', '说明'], env_rows)

doc.add_heading('8.2 服务启动说明', level=2)
doc.add_paragraph('各服务的启动方式：')

doc.add_heading('1) 数据库初始化', level=3)
doc.add_paragraph('使用 mydb.sql 脚本初始化 MySQL 数据库：')
p = doc.add_paragraph('mysql -u root -p mydb < mydb.sql')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(11)

doc.add_heading('2) Java 后端启动', level=3)
doc.add_paragraph('进入 Edu_platform 目录，使用 Maven 启动 Spring Boot：')
p = doc.add_paragraph('cd Edu_platform && mvn spring-boot:run')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(11)

doc.add_heading('3) Python AI 服务启动', level=3)
doc.add_paragraph('进入 Edu_py 目录，启动 FastAPI 服务：')
p = doc.add_paragraph('cd Edu_py && python main.py\n# 或使用 uvicorn：\nuvicorn main:app --host 0.0.0.0 --port 8000')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(11)

doc.add_heading('4) 前端启动', level=3)
doc.add_paragraph('进入 EduGenius 目录，安装依赖并启动开发服务器：')
p = doc.add_paragraph('cd EduGenius && npm install && npm run dev')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(11)

doc.add_heading('5) 数字人服务启动', level=3)
doc.add_paragraph('进入 avatar_py 目录，启动数字人服务：')
p = doc.add_paragraph('cd avatar_py && python app.py --model musetalk --transport webrtc --listenport 8010')
for r in p.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(11)

doc.add_heading('6) 一键启动', level=3)
doc.add_paragraph('项目根目录提供 quick_start.ps1 PowerShell 脚本，可一键启动所有服务。')

doc.add_page_break()

# ==================== 9. 总结 ====================
doc.add_heading('9. 总结', level=1)
doc.add_paragraph(
    'EduGenius 智能教育平台采用多服务协同的微服务架构，将传统业务逻辑（Java Spring Boot）'
    '与 AI 智能服务（Python FastAPI）分离，前端使用 Vue 3 + PrimeVue 构建现代化交互界面，'
    '并集成了基于 WebRTC 的数字人交互系统。'
)
doc.add_paragraph(
    '系统的核心创新点在于：'
)
innovations = [
    'Agent 智能体架构：7个专职 Agent 覆盖教学全流程，通过统一 LLM Adapter 实现多模型切换',
    '知识图谱驱动：基于学生答题记录自动构建个人知识图谱，实现精准学情分析',
    '试卷自动化处理：上传试卷即可自动生成标准答案、评分标准和试卷分析报告',
    '数字人实时交互：基于 MuseTalk 的唇形同步技术，提供沉浸式 AI 对话体验',
    '多角色协同：教师、学生、管理员三种角色，各自拥有专属功能模块',
]
for inn in innovations:
    add_bullet(inn)

doc.add_paragraph()
doc.add_paragraph(
    '整个系统具有良好的模块化设计和可扩展性，各服务独立部署、松耦合，'
    '便于后续功能迭代和技术升级。'
)

# ========== 保存文档 ==========
output_path = r'C:\Users\lenovo\IdeaProjects\01014097\EduGenius_技术文档.docx'
doc.save(output_path)
print(f'✅ 技术文档已生成：{output_path}')
